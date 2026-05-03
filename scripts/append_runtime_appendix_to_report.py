#!/usr/bin/env python3
"""Append local runtime screenshots to the Final Project Report docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs/Final_Project_Report_Intelligent_Learning_Assistant_LightRAG_Agent_updated.docx"
ATTACH_DIR = ROOT / "docs/report_runtime_attachments"

SECTION_TITLE = "附录 · 运行时界面截图（本地启动实测）"

FIGURES: list[tuple[str, str]] = [
    (
        "ss_01_studyforge_home.png",
        "图附-1 StudyForge 首页 — Knowledge Bases 栏目（尚无 Subject）",
    ),
    (
        "ss_02_studyforge_swagger.png",
        "图附-2 StudyForge 后端 Swagger UI（http://127.0.0.1:8010/docs）",
    ),
    (
        "ss_03_create_subject_dialog.png",
        "图附-3 新建 Subject 对话框",
    ),
    (
        "ss_04_subject_workspace.png",
        "图附-4 Subject 工作区 — 讲义与试卷上传入口（示例 Subject：Course Demo）",
    ),
    (
        "ss_05_chat_documents_panel.png",
        "图附-5 对话视图 — 提问输入框与右侧 Documents / Mind Map 面板",
    ),
    (
        "ss_06_api_root_fabricated.png",
        "图附-6 API 根路径 JSON（Chrome 对 JSON 页渲染不适配截图时，使用离线排版示意）",
    ),
]


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Missing report: {DOC_PATH}")

    doc = Document(str(DOC_PATH))

    if any(p.text.strip() == SECTION_TITLE for p in doc.paragraphs):
        raise SystemExit("Appendix section already present; delete manually before re-running.")

    doc.add_page_break()
    doc.add_heading(SECTION_TITLE, level=1)
    intro = (
        "以下为在本机启动前后端进行冒烟测试时的界面截图，用于报告附件。"
        "测试配置：前端 Vite http://127.0.0.1:5173 ，环境变量 "
        "VITE_DEV_API_PROXY_TARGET=http://127.0.0.1:8010 ；后端 uvicorn StudyForge "
        "http://127.0.0.1:8010 。"
        "仓库：https://github.com/arthurpanhku/studyforge"
    )
    doc.add_paragraph(intro)

    for fname, caption in FIGURES:
        path = ATTACH_DIR / fname
        if not path.exists():
            raise SystemExit(f"Missing attachment image: {path}")
        doc.add_paragraph(caption)
        doc.add_picture(str(path), width=Inches(6.5))
        doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

    doc.save(str(DOC_PATH))
    print(f"Updated {DOC_PATH}")


if __name__ == "__main__":
    main()
