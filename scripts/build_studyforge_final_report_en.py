#!/usr/bin/env python3
"""Build a clean English Final Project Report (.docx) for StudyForge with diagrams + screenshots."""
from __future__ import annotations

import subprocess
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_SCRIPTS_DIR = Path(__file__).resolve().parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))
from studyforge_report_en_body import add_long_form_sections  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
OUT_DOC = ROOT / "docs/StudyForge_Final_Project_Report_EN.docx"
ASSETS = ROOT / "docs/report_en_assets"
LEGACY_SHOTS = ROOT / "docs/report_runtime_attachments"

REPO = "https://github.com/arthurpanhku/studyforge"


def ensure_assets() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)


def shade_header_cells(table) -> None:
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF7")
        tc_pr.append(shd)


def set_doc_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        hs = doc.styles[name]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)
    for sec in doc.sections:
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("StudyForge")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(
        "Intelligent Learning Assistant — Knowledge-Graph RAG & Agent Orchestration\nFinal Project Report"
    )
    rs.font.size = Pt(14)
    rs.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt.add_run("May 2026").italic = True

    doc.add_paragraph()


def add_team_table(doc: Document) -> None:
    cap = doc.add_paragraph("Team Members")
    cap.runs[0].bold = True
    cap.runs[0].font.size = Pt(12)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ("Name", "University Number", "HKU Email"),
        ("Pan Chao", "3036383769", "u3638376@connect.hku.hk"),
        ("Liu Xunyu", "3036656441", "u3665644@connect.hku.hk"),
        ("Huang Zhenxiang", "3036655631", "u3665563@connect.hku.hk"),
        ("Zhang Yucheng", "3036658671", "u3665867@connect.hku.hk"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                if i == 0:
                    for run in paragraph.runs:
                        run.bold = True
    shade_header_cells(table)
    doc.add_paragraph()


def h(doc: Document, level: int, text: str) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)


def body(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        if not block.strip():
            continue
        p = doc.add_paragraph(block.strip())
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15


def bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item.strip(), style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15


def add_picture_caption(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(10)
    cr = cp.add_run(caption)
    cr.bold = True
    cr.font.size = Pt(10)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].paragraph_format.space_after = Pt(14)


def dot_png(dot: str, out: Path, target_wh: tuple[int, int]) -> None:
    proc = subprocess.run(
        ["dot", "-Tpng", "-o", str(out)],
        input=dot.encode("utf-8"),
        capture_output=True,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode("utf-8", errors="replace"))

    from PIL import Image

    im = Image.open(out).convert("RGBA")
    im = im.resize(target_wh, Image.Resampling.LANCZOS)
    bg = Image.new("RGB", target_wh, (255, 255, 255))
    bg.paste(im, mask=im.split()[3])
    bg.save(out, format="PNG", optimize=True)


def gen_architecture_diagram() -> Path:
    path = ASSETS / "fig_architecture_en.png"
    dot = rf"""
digraph G {{
  graph [fontname="Helvetica", fontsize=11, bgcolor=white, pad=0.35, rankdir=TB];
  node [fontname="Helvetica", fontsize=10, shape=box, style="rounded,filled",
        fillcolor="#E8F4FF", color="#3498DB"];
  edge [fontname="Helvetica", fontsize=9, color="#555555"];

  labelloc="t";
  label="StudyForge — layered architecture\\nRepository: {REPO}";

  FE [label="Web client\\nVue 3 · Vite · Element Plus"];
  API [label="API gateway\\nFastAPI · streaming NDJSON · tools"];
  LKG [label="LightRAG core\\nchunking · embeddings · KG + vectors"];
  AG [label="Agents\\nregistered tools · citations · variants"];
  LLM [label="Configured LLMs\\nextraction · chat · grading"];

  FE -> API -> LKG -> LLM;
  API -> AG -> LKG [style=dashed, label="tool calls"];
}}
"""
    dot_png(dot, path, (1100, 560))
    return path


def gen_mock_rag_panel_png() -> Path:
    """Simulated assistant reply when live LLM keys are absent — illustrative only."""
    from PIL import Image, ImageDraw, ImageFont

    path = ASSETS / "fig_mock_rag_response_en.png"
    W, H = 980, 520
    img = Image.new("RGB", (W, H), (250, 251, 252))
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 18)
        body_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 14)
        small = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 12)
    except OSError:
        title_font = body_font = small = ImageFont.load_default()

    draw.rounded_rectangle((24, 24, W - 24, H - 24), radius=12, outline="#CBD5E1", width=2)
    draw.text((44, 44), "Simulated assistant reply (offline demonstration)", fill="#1E293B", font=title_font)

    sample = (
        "Based on the uploaded syllabus excerpt (mock retrieval):\n"
        "• Core topic: structured retrieval with LightRAG hybrid modes (local / global / mix).\n"
        "• Evidence tie-in: chunk #3 — definitions of entities and relations before graph fusion.\n"
        "• References: [[mock-file-id|p.2]] (placeholder anchors for traceability).\n\n"
        "Note: This panel is generated for reporting when production LLM / embedding keys\n"
        "are not configured; live answers follow the same schema once indexing succeeds."
    )
    y = 88
    for line in sample.split("\n"):
        draw.text((44, y), line, fill="#334155", font=body_font)
        y += 22

    draw.text((44, H - 52), "StudyForge · mock UI artifact — not from a live model call", fill="#94A3B8", font=small)
    img.save(path)
    return path


def gen_settings_mock_png() -> Path:
    """Simulated Settings panel — English labels when credentials are omitted."""
    from PIL import Image, ImageDraw, ImageFont

    path = ASSETS / "fig_settings_mock_en.png"
    W, H = 720, 540
    img = Image.new("RGB", (W, H), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        title_f = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 20)
        lab_f = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 14)
        hint_f = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 12)
    except OSError:
        title_f = lab_f = hint_f = ImageFont.load_default()

    draw.rectangle((0, 0, W, 52), fill="#B8775F")
    draw.text((24, 14), "Settings — LLM configuration (mock)", fill="white", font=title_f)

    fields = [
        ("Knowledge graph (extraction)", "(API key not set — indexing stops until configured)"),
        ("Chat (query)", "(optional override for answering)"),
        ("Embedding", "(required for vector index)"),
    ]
    y = 72
    for label, hint in fields:
        draw.text((32, y), label, fill="#111827", font=lab_f)
        draw.rounded_rectangle((32, y + 26, W - 32, y + 52), radius=4, outline="#CBD5E1", width=2)
        draw.text((42, y + 32), hint, fill="#94A3B8", font=hint_f)
        y += 96

    draw.text(
        (32, H - 72),
        "Illustrative mock for documentation when keys are omitted.\n"
        "Live deployments bind keys via the in-app Settings dialog or server-side secrets.",
        fill="#64748B",
        font=hint_f,
    )
    img.save(path)
    return path


def copy_if_exists(name: str) -> Path | None:
    src = LEGACY_SHOTS / name
    if not src.exists():
        return None
    dst = ASSETS / name
    dst.write_bytes(src.read_bytes())
    return dst


def main() -> None:
    ensure_assets()
    arch = gen_architecture_diagram()
    mock_panel = gen_mock_rag_panel_png()
    settings_mock = gen_settings_mock_png()

    doc = Document()
    set_doc_styles(doc)

    add_title_block(doc)
    add_team_table(doc)
    doc.add_page_break()

    add_long_form_sections(
        doc,
        h=h,
        body=body,
        bullet_list=bullet_list,
        after_section3_heading=lambda d: add_picture_caption(
            d, arch, "Figure 1. Layered architecture of StudyForge (diagram)."
        ),
    )

    h(doc, 1, "6. Figures — Screenshots & Illustrative Panels")
    shots = [
        ("ss_01_studyforge_home.png", "Figure 2. Home — knowledge bases overview."),
        ("ss_02_studyforge_swagger.png", "Figure 3. OpenAPI / Swagger UI listing StudyForge routes."),
        ("ss_03_create_subject_dialog.png", "Figure 4. Creating a subject (knowledge base)."),
        ("ss_04_subject_workspace.png", "Figure 5. Subject workspace — courseware & exam upload tiles."),
        ("ss_05_chat_documents_panel.png", "Figure 6. Chat layout with documents side panel."),
        ("ss_06_api_root_fabricated.png", "Figure 7. Root endpoint payload snapshot (typed reconstruction)."),
    ]
    for fname, caption in shots:
        copied = copy_if_exists(fname)
        if copied:
            add_picture_caption(doc, copied, caption)
        else:
            doc.add_paragraph(f"[Missing asset: {fname}]")

    add_picture_caption(doc, mock_panel, "Figure 8. Mock grounded reply template (offline illustration).")

    add_picture_caption(
        doc,
        settings_mock,
        "Figure 9. Mock Settings panel — credential slots required before live LightRAG indexing.",
        width_in=6.0,
    )

    body(
        doc,
        (
            "Figures 10–11 are refreshed via `scripts/capture_report_figures_10_11.py` against a running stack: chat "
            "and knowledge-graph extraction use DeepSeek’s OpenAI-compatible API (`https://api.deepseek.com`, see "
            "https://api-docs.deepseek.com/), while embeddings for LightRAG indexing typically remain on a dedicated "
            "embedding provider (e.g., SiliconFlow) unless locally hosted alternatives are configured. Figure 10 "
            "captures the chat workspace with the Documents-side panel opened into the knowledge-graph viewer; "
            "Figure 11 renders structured variant items produced by the `/variant-questions` endpoint bound to mixed "
            "retrieval context."
        ),
    )
    rag_kg_shot = ASSETS / "user_feat_rag_knowledge_graph.png"
    exam_gen_shot = ASSETS / "user_feat_agent_exam_generation.png"
    if rag_kg_shot.exists():
        add_picture_caption(
            doc,
            rag_kg_shot,
            "Figure 10. Chat + Documents panel with knowledge-graph dialog (DeepSeek-backed chat/KG after indexing).",
            width_in=6.2,
        )
    else:
        doc.add_paragraph("[Missing asset: user_feat_rag_knowledge_graph.png]")
    if exam_gen_shot.exists():
        add_picture_caption(
            doc,
            exam_gen_shot,
            "Figure 11. RAG-bound variant questions (mix retrieval + DeepSeek chat per api-docs.deepseek.com).",
            width_in=6.2,
        )
    else:
        doc.add_paragraph("[Missing asset: user_feat_agent_exam_generation.png]")

    body(
        doc,
        (
            "Indexing still requires a working embedding provider (SiliconFlow, OpenAI-compatible embeddings, or local "
            "Ollama) in addition to chat/extraction keys; without embeddings, document status remains incomplete even "
            "when chat completes successfully."
        ),
    )

    h(doc, 1, "7. Conclusion and Outlook")
    body(
        doc,
        (
            "StudyForge situates contemporary foundation models inside a retrieval discipline appropriate for education: "
            "evidence-first answering, inspectable graph structure, and workflows that reuse shared primitives across chat, "
            "visualization, and exam analytics. Continued work belongs where deployments always do—latency-aware indexing, "
            "richer evaluation harnesses tied to course outcomes, and tighter governance patterns for tenant isolation "
            "when institutions adopt the stack beyond single-subject pilots."
        ),
    )

    h(doc, 1, "References")
    body(doc, f"StudyForge source repository: {REPO}")
    body(
        doc,
        (
            "DeepSeek API documentation (OpenAI-compatible integration): https://api-docs.deepseek.com/"
        ),
    )
    body(
        doc,
        (
            "HKU COMP7607 course materials (internal). LightRAG upstream MIT-licensed subtree vendored under `/LightRAG`. "
            "Further bibliography entries may be appended per departmental citation style."
        ),
    )

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOC))
    print(f"Wrote {OUT_DOC}")


if __name__ == "__main__":
    main()
